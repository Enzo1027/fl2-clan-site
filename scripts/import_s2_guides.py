from __future__ import annotations

import json
import shutil
import subprocess
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path("/Users/davidprice/Downloads/S2 Guide")
PUBLIC = ROOT / "public"
ASSET_ROOT = PUBLIC / "assets" / "guides" / "s2"
UPLOAD_ROOT = PUBLIC / "assets" / "uploads"
DATA_DIR = PUBLIC / "data"
SOURCE_DOCS = ROOT / "source-documents" / "s2"
MASTER_GUIDE_ROOT = SOURCE / "Server 630 - Season 2 Everfrost Master Guide"


LANGUAGES = [
    {"code": "en", "name": "English", "native": "English", "short": "EN", "flag": "🇺🇸", "dir": "ltr"},
    {"code": "es", "name": "Spanish", "native": "Español", "short": "ES", "flag": "🇪🇸", "dir": "ltr"},
    {"code": "fr", "name": "French", "native": "Français", "short": "FR", "flag": "🇫🇷", "dir": "ltr"},
    {"code": "de", "name": "German", "native": "Deutsch", "short": "DE", "flag": "🇩🇪", "dir": "ltr"},
    {"code": "ar", "name": "Arabic", "native": "العربية", "short": "AR", "flag": "🇸🇦", "dir": "rtl"},
    {"code": "tr", "name": "Turkish", "native": "Türkçe", "short": "TR", "flag": "🇹🇷", "dir": "ltr"},
    {"code": "nl", "name": "Dutch", "native": "Nederlands", "short": "NL", "flag": "🇳🇱", "dir": "ltr"},
    {"code": "id", "name": "Indonesian", "native": "Bahasa Indonesia", "short": "ID", "flag": "🇮🇩", "dir": "ltr"},
]

LANGUAGE_ALIASES = {
    "en": ["english", " z day"],
    "es": ["spanish"],
    "fr": ["french"],
    "de": ["german"],
    "ar": ["arabic"],
    "tr": ["turkish"],
    "nl": [" nl ", " dutch ", " nederlands "],
    "id": ["indonesian", " indonesia ", " bahasa indonesia "],
}

TOTAL_DAYS = 7


IMAGE_SOURCES = {
    1: {
        "en": "z day1.png",
        "es": "Spanish Day 1.png",
        "fr": "French Day 1.png",
        "de": "German.png",
        "ar": "Arabic Day 1.png",
        "tr": "Turkish.png",
        "nl": "NL.png",
        "id": "Indonesian.png",
    },
    2: {
        "en": "S2 Day 2 English.png",
        "es": "S2 Day 2 Spanish.png",
        "fr": "S2 Day 2 French.png",
        "de": "S2 Day 2 German.png",
        "ar": "S2 Day 2 Arabic.png",
        "tr": "S2 Day 2 Turkish.png",
        "nl": "S2 Day 2 NL.png",
        "id": "S2 Day 2 Indonesian.png",
    },
    3: {
        "en": "S2 Day 3 English.png",
        "es": "S2 Day 3 Spanish.png",
        "fr": "S2 Day 3 French.png",
        "de": "S2 Day 3 German.png",
        "ar": "S2 Day 3 Arabic.png",
        "tr": "S2 Day 3 Turkish.png",
        "nl": "S2 Day 3 NL.png",
        "id": "S2 Day 3 Indonesian.png",
    },
    4: {
        "en": "S2 Day 4 English.png",
        "es": "S2 Day 4 Spanish.png",
        "fr": "S2 Day 4 French.png",
        "de": "S2 Day 4 German.png",
        "ar": "S2 Day 4 Arabic.png",
        "tr": "S2 Day 4 Turkish.png",
        "nl": "S2 Day 4 NL.png",
        "id": "S2 Day 4 Indonesian.png",
    },
    5: {
        "en": "S2 Day 5 English.png",
        "es": "S2 Day 5 Spanish.png",
        "fr": "S2 Day 5 French.png",
        "de": "S2 Day 5 German.png",
        "ar": "S2 Day 5 Arabic.png",
        "tr": "S2 Day 5 Turkish.png",
        "nl": "S2 Day 5 NL.png",
        "id": "S2 Day 5 Indonesian.png",
    },
    6: {
        "en": "S2 Day 6 English.png",
        "es": "S2 Day 6 Spanish.png",
        "fr": "S2 Day 6 French.png",
        "de": "S2 Day 6 German.png",
        "ar": "S2 Day 6 Arabic.png",
        "tr": "S2 Day 6 Turkish.png",
        "nl": "S2 Day 6 NL.png",
        "id": "S2 Day 6 Indonesian.png",
    },
    7: {
        "en": "S2 Day 7 English.png",
        "es": "S2 Day 7 Spanish.png",
        "fr": "S2 Day 7 French.png",
        "de": "S2 Day 7 German.png",
        "ar": "S2 Day 7 Arabic.png",
        "tr": "S2 Day 7 Turkish.png",
        "nl": "S2 Day 7 NL.png",
        "id": "S2 Day 7 Indonesian.png",
    },
}


DOCX_SOURCES = {
    1: [
        "FL2_Day1_Alliance_Mail_With_Icons.docx",
        "FL2_Day1_Master_Plan_Word_Outputs_DE_ES_AR_FR.docx",
        "FL2_Day1_Master_Plan_NL.docx",
        "FL2_Day1_Master_Plan_Indonesian.docx",
    ],
    2: [
        "S2 Day 2 English.docx",
        "FL2_Day2_Master_Plan_All_Languages.docx",
        "FL2_Day2_Master_Plan_NL.docx",
        "FL2_Day2_Master_Plan_Indonesian.docx",
    ],
    3: [
        "S2 Day 3 English.docx",
        "FL2_Day3_Master_Plan_All_Languages.docx",
        "FL2_Day3_Master_Plan_NL.docx",
        "FL2_Day3_Master_Plan_Indonesian.docx",
    ],
    4: [
        "FL2_Day4_Alliance_Mail_With_Icons.docx",
        "FL2_Day4_Master_Plan_All_Languages.docx",
        "FL2_Day4_Master_Plan_NL.docx",
        "FL2_Day4_Master_Plan_Indonesian.docx",
    ],
    5: [
        "FL2_Day5_Alliance_Mail_With_Icons.docx",
        "FL2_Day5_Master_Plan_All_Languages.docx",
    ],
    6: [
        "FL2_Day6_Alliance_Mail_With_Icons.docx",
        "FL2_Day6_Master_Plan_All_Languages.docx",
    ],
    7: [
        "FL2_Day7_Alliance_Mail_With_Icons.docx",
        "FL2_Day7_Master_Plan_All_Languages.docx",
    ],
}

MASTER_GUIDE_ID = "s2-everfrost-master-guide"
MASTER_GUIDE_TITLE = "Server 630 Season 2 Everfrost Master Guide"
MASTER_GUIDE_SOURCES = {
    "en": MASTER_GUIDE_ROOT / "_EN - ORGINAL" / "Server_630_Season_2_Everfrost_Master_Guide_EN.pdf",
    "es": MASTER_GUIDE_ROOT / "ES" / "Server_630_Season_2_Everfrost_Master_Guide_ES.pdf",
    "fr": MASTER_GUIDE_ROOT / "FR" / "Server_630_Season_2_Everfrost_Master_Guide_FR.pdf",
    "de": MASTER_GUIDE_ROOT / "DE" / "Server_630_Season_2_Everfrost_Master_Guide_DE.pdf",
    "ar": MASTER_GUIDE_ROOT / "AR" / "Server_630_Season_2_Everfrost_Master_Guide_AR.pdf",
    "tr": MASTER_GUIDE_ROOT / "TR" / "Server_630_Season_2_Everfrost_Master_Guide_TR.pdf",
    "nl": MASTER_GUIDE_ROOT / "DU" / "Server_630_Season_2_Everfrost_Master_Guide_NL.pdf",
}


@dataclass
class BlockRange:
    code: str
    start: int
    end: int | None


DAY1_PARAGRAPH_RANGES = [
    BlockRange("de", 0, 101),
    BlockRange("es", 101, 202),
    BlockRange("ar", 202, 303),
    BlockRange("fr", 303, None),
]

DAY3_BLOCK_RANGES = [
    BlockRange("ar", 0, 34),
    BlockRange("de", 34, 68),
    BlockRange("es", 68, 102),
    BlockRange("fr", 102, 136),
    BlockRange("tr", 136, None),
]

DAY4_BLOCK_RANGES = [
    BlockRange("ar", 0, 99),
    BlockRange("de", 99, 197),
    BlockRange("es", 197, 295),
    BlockRange("fr", 295, 393),
    BlockRange("tr", 393, None),
]


def child_to_block(child, document: Document) -> dict | None:
    if isinstance(child, CT_P):
        paragraph = Paragraph(child, document)
        text = paragraph.text.strip()
        if not text:
            return None
        return {"type": "paragraph", "text": text}

    if isinstance(child, CT_Tbl):
        table = Table(child, document)
        rows = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                row_cells.append(lines)
            rows.append(row_cells)
        return {"type": "table", "rows": rows}

    return None


def iter_blocks(path: Path) -> list[dict]:
    document = Document(path)
    blocks = []
    for child in document.element.body.iterchildren():
        block = child_to_block(child, document)
        if block:
            blocks.append(block)
    return blocks


def paragraph_blocks(path: Path) -> list[dict]:
    document = Document(path)
    return [
        {"type": "paragraph", "text": paragraph.text.strip()}
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]


def split_ranges(blocks: list[dict], ranges: list[BlockRange]) -> dict[str, list[dict]]:
    split = {}
    for item in ranges:
        split[item.code] = blocks[item.start : item.end]
    return split


def split_labeled_sections(blocks: list[dict]) -> dict[str, list[dict]]:
    labels = {
        "ARABIC": "ar",
        "Arabic / العربية": "ar",
        "GERMAN": "de",
        "German / Deutsch": "de",
        "SPANISH": "es",
        "Spanish / Español": "es",
        "FRENCH": "fr",
        "French / Français": "fr",
        "TURKISH": "tr",
        "Turkish / Türkçe": "tr",
        "Dutch / Nederlands": "nl",
        "Indonesian / Bahasa Indonesia": "id",
    }
    sections: dict[str, list[dict]] = {}
    current = None
    for block in blocks:
        text = block.get("text", "")
        if block["type"] == "paragraph" and text in labels:
            current = labels[text]
            sections[current] = []
            continue
        if current:
            sections[current].append(block)
    return sections


def split_fixed_sections(blocks: list[dict], codes: list[str]) -> dict[str, list[dict]]:
    if not codes:
        return {}
    section_size = len(blocks) // len(codes)
    sections = {}
    for index, code in enumerate(codes):
        start = index * section_size
        end = (index + 1) * section_size if index < len(codes) - 1 else None
        sections[code] = blocks[start:end]
    return sections


def copy_assets() -> list[dict]:
    days = []
    for day in range(1, TOTAL_DAYS + 1):
        language_files = IMAGE_SOURCES.get(day, {})
        source_day = SOURCE / f"S2 Day {day}"
        day_exists = source_day.is_dir()
        entries = {}
        if day_exists:
            target_day = ASSET_ROOT / f"day-{day}"
            target_day.mkdir(parents=True, exist_ok=True)
            if not language_files:
                language_files = discover_images(source_day)

            for code, filename in language_files.items():
                source_path = source_day / filename
                if not source_path.exists():
                    continue
                target_path = target_day / f"{code}.png"
                shutil.copy2(source_path, target_path)
                entries[code] = {
                    "image": f"assets/guides/s2/day-{day}/{code}.png",
                    "sourceFile": filename,
                }

            source_doc_day = SOURCE_DOCS / f"day-{day}"
            source_doc_day.mkdir(parents=True, exist_ok=True)
            for filename in DOCX_SOURCES.get(day, []):
                source_doc = source_day / filename
                if source_doc.exists():
                    shutil.copy2(source_doc, source_doc_day / filename)

        days.append(
            {
                "id": f"day-{day}",
                "number": day,
                "title": f"S2 Day {day}",
                "label": f"Day {day}",
                "season": "Season 2",
                "status": "live" if entries else "upcoming",
                "languages": entries,
            }
        )
    return days


def discover_images(source_day: Path) -> dict[str, str]:
    discovered = {}
    images = [path for path in source_day.iterdir() if path.suffix.lower() == ".png"]
    for code, aliases in LANGUAGE_ALIASES.items():
        for path in images:
            name = f" {path.stem.lower()} "
            if any(alias in name for alias in aliases):
                discovered[code] = path.name
                break
    return discovered


def extract_docs() -> dict:
    docs: dict[str, dict[str, dict]] = {f"day-{day}": {} for day in range(1, TOTAL_DAYS + 1)}

    day1_en = SOURCE / "S2 Day 1" / "FL2_Day1_Alliance_Mail_With_Icons.docx"
    if day1_en.exists():
        docs["day-1"]["en"] = {
            "source": day1_en.name,
            "title": "S2 Day 1 Master Plan",
            "blocks": iter_blocks(day1_en),
        }

    day1 = SOURCE / "S2 Day 1" / "FL2_Day1_Master_Plan_Word_Outputs_DE_ES_AR_FR.docx"
    for code, blocks in split_ranges(paragraph_blocks(day1), DAY1_PARAGRAPH_RANGES).items():
        docs["day-1"][code] = {
            "source": day1.name,
            "title": "S2 Day 1 Master Plan",
            "blocks": blocks,
        }

    day1_nl = SOURCE / "S2 Day 1" / "FL2_Day1_Master_Plan_NL.docx"
    if day1_nl.exists():
        docs["day-1"]["nl"] = {
            "source": day1_nl.name,
            "title": "S2 Day 1 Master Plan",
            "blocks": iter_blocks(day1_nl),
        }

    day1_id = SOURCE / "S2 Day 1" / "FL2_Day1_Master_Plan_Indonesian.docx"
    if day1_id.exists():
        docs["day-1"]["id"] = {
            "source": day1_id.name,
            "title": "S2 Day 1 Master Plan",
            "blocks": iter_blocks(day1_id),
        }

    day2_en = SOURCE / "S2 Day 2" / "S2 Day 2 English.docx"
    docs["day-2"]["en"] = {
        "source": day2_en.name,
        "title": "S2 Day 2 Alliance Mails",
        "blocks": iter_blocks(day2_en),
    }

    day2_all = SOURCE / "S2 Day 2" / "FL2_Day2_Master_Plan_All_Languages.docx"
    for code, blocks in split_labeled_sections(iter_blocks(day2_all)).items():
        docs["day-2"][code] = {
            "source": day2_all.name,
            "title": "S2 Day 2 Alliance Mails",
            "blocks": blocks,
        }

    day2_nl = SOURCE / "S2 Day 2" / "FL2_Day2_Master_Plan_NL.docx"
    if day2_nl.exists():
        docs["day-2"]["nl"] = {
            "source": day2_nl.name,
            "title": "S2 Day 2 Master Plan",
            "blocks": iter_blocks(day2_nl),
        }

    day2_id = SOURCE / "S2 Day 2" / "FL2_Day2_Master_Plan_Indonesian.docx"
    if day2_id.exists():
        docs["day-2"]["id"] = {
            "source": day2_id.name,
            "title": "S2 Day 2 Master Plan",
            "blocks": iter_blocks(day2_id),
        }

    day3_en = SOURCE / "S2 Day 3" / "S2 Day 3 English.docx"
    docs["day-3"]["en"] = {
        "source": day3_en.name,
        "title": "S2 Day 3 Master Plan",
        "blocks": iter_blocks(day3_en),
    }

    day3_all = SOURCE / "S2 Day 3" / "FL2_Day3_Master_Plan_All_Languages.docx"
    for code, blocks in split_ranges(iter_blocks(day3_all), DAY3_BLOCK_RANGES).items():
        docs["day-3"][code] = {
            "source": day3_all.name,
            "title": "S2 Day 3 Master Plan",
            "blocks": blocks,
        }

    day3_nl = SOURCE / "S2 Day 3" / "FL2_Day3_Master_Plan_NL.docx"
    if day3_nl.exists():
        docs["day-3"]["nl"] = {
            "source": day3_nl.name,
            "title": "S2 Day 3 Master Plan",
            "blocks": iter_blocks(day3_nl),
        }

    day3_id = SOURCE / "S2 Day 3" / "FL2_Day3_Master_Plan_Indonesian.docx"
    if day3_id.exists():
        docs["day-3"]["id"] = {
            "source": day3_id.name,
            "title": "S2 Day 3 Master Plan",
            "blocks": iter_blocks(day3_id),
        }

    day4_en = SOURCE / "S2 Day 4" / "FL2_Day4_Alliance_Mail_With_Icons.docx"
    if day4_en.exists():
        docs["day-4"]["en"] = {
            "source": day4_en.name,
            "title": "S2 Day 4 Master Plan",
            "blocks": iter_blocks(day4_en),
        }

    day4_all = SOURCE / "S2 Day 4" / "FL2_Day4_Master_Plan_All_Languages.docx"
    if day4_all.exists():
        for code, blocks in split_ranges(iter_blocks(day4_all), DAY4_BLOCK_RANGES).items():
            docs["day-4"][code] = {
                "source": day4_all.name,
                "title": "S2 Day 4 Master Plan",
                "blocks": blocks,
            }

    day4_nl = SOURCE / "S2 Day 4" / "FL2_Day4_Master_Plan_NL.docx"
    if day4_nl.exists():
        docs["day-4"]["nl"] = {
            "source": day4_nl.name,
            "title": "S2 Day 4 Master Plan",
            "blocks": iter_blocks(day4_nl),
        }

    day4_id = SOURCE / "S2 Day 4" / "FL2_Day4_Master_Plan_Indonesian.docx"
    if day4_id.exists():
        docs["day-4"]["id"] = {
            "source": day4_id.name,
            "title": "S2 Day 4 Master Plan",
            "blocks": iter_blocks(day4_id),
        }

    day5_en = SOURCE / "S2 Day 5" / "FL2_Day5_Alliance_Mail_With_Icons.docx"
    if day5_en.exists():
        docs["day-5"]["en"] = {
            "source": day5_en.name,
            "title": "S2 Day 5 Master Plan",
            "blocks": iter_blocks(day5_en),
        }

    day5_all = SOURCE / "S2 Day 5" / "FL2_Day5_Master_Plan_All_Languages.docx"
    if day5_all.exists():
        for code, blocks in split_labeled_sections(iter_blocks(day5_all)).items():
            docs["day-5"][code] = {
                "source": day5_all.name,
                "title": "S2 Day 5 Master Plan",
                "blocks": blocks,
            }

    day6_en = SOURCE / "S2 Day 6" / "FL2_Day6_Alliance_Mail_With_Icons.docx"
    if day6_en.exists():
        docs["day-6"]["en"] = {
            "source": day6_en.name,
            "title": "S2 Day 6 Master Plan",
            "blocks": iter_blocks(day6_en),
        }

    day6_all = SOURCE / "S2 Day 6" / "FL2_Day6_Master_Plan_All_Languages.docx"
    if day6_all.exists():
        for code, blocks in split_fixed_sections(
            iter_blocks(day6_all),
            ["ar", "de", "es", "fr", "tr", "nl", "id"],
        ).items():
            docs["day-6"][code] = {
                "source": day6_all.name,
                "title": "S2 Day 6 Master Plan",
                "blocks": blocks,
            }

    day7_en = SOURCE / "S2 Day 7" / "FL2_Day7_Alliance_Mail_With_Icons.docx"
    if day7_en.exists():
        docs["day-7"]["en"] = {
            "source": day7_en.name,
            "title": "S2 Day 7 Master Plan",
            "blocks": iter_blocks(day7_en),
        }

    day7_all = SOURCE / "S2 Day 7" / "FL2_Day7_Master_Plan_All_Languages.docx"
    if day7_all.exists():
        for code, blocks in split_labeled_sections(iter_blocks(day7_all)).items():
            docs["day-7"][code] = {
                "source": day7_all.name,
                "title": "S2 Day 7 Master Plan",
                "blocks": blocks,
            }

    return docs


def page_sort_key(path: Path) -> int:
    match = re.search(r"-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else 0


def render_pdf_pages(source: Path, pages_dir: Path) -> list[dict]:
    pages_dir.mkdir(parents=True, exist_ok=True)
    for path in pages_dir.glob("*.png"):
        path.unlink()

    prefix = pages_dir / "render"
    subprocess.run(["pdftoppm", "-r", "160", "-png", str(source), str(prefix)], check=True)

    pages = []
    rendered = sorted(pages_dir.glob("render-*.png"), key=page_sort_key)
    for index, path in enumerate(rendered, start=1):
        target = pages_dir / f"page-{index:02d}.png"
        path.rename(target)
        pages.append(
            {
                "number": index,
                "image": str(target.relative_to(PUBLIC)),
            }
        )
    return pages


def build_uploads() -> list[dict]:
    target_dir = UPLOAD_ROOT / MASTER_GUIDE_ID
    if target_dir.exists():
        shutil.rmtree(target_dir)

    languages = {}
    for code, source in MASTER_GUIDE_SOURCES.items():
        if code == "en" and not source.exists():
            source = SOURCE / "Server_630_Season_2_Everfrost_Master_Guide_EN.pdf"
        if not source.exists():
            continue

        target_lang_dir = target_dir / code
        target_lang_dir.mkdir(parents=True, exist_ok=True)
        pdf_target = target_lang_dir / source.name
        shutil.copy2(source, pdf_target)
        pages = render_pdf_pages(source, target_lang_dir / "pages")
        languages[code] = {
            "href": str(pdf_target.relative_to(PUBLIC)),
            "pageCount": len(pages),
            "pages": pages,
            "sourceFile": source.name,
        }

    if not languages:
        return []

    return [
        {
            "id": MASTER_GUIDE_ID,
            "title": MASTER_GUIDE_TITLE,
            "type": "PDF",
            "status": "Live",
            "featured": True,
            "languages": languages,
        }
    ]


def write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    manifest = {
        "clan": "FrostBorn Lions",
        "tag": "FL2",
        "season": "Season 2",
        "languages": LANGUAGES,
        "days": copy_assets(),
        "uploads": build_uploads(),
        "library": [
            {"id": "s2-guides", "title": "Season 2 Guides", "status": "live"},
            {"id": "alliance-mails", "title": "Alliance Mail Text", "status": "live"},
            {"id": "future-uploads", "title": "Upload Queue", "status": "ready"},
        ],
    }
    write_json(DATA_DIR / "manifest.json", manifest)
    write_json(DATA_DIR / "docs.json", extract_docs())


if __name__ == "__main__":
    main()
